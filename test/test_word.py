
from docx import Document
import openxmllib

doc = openxmllib.openXmlDocument(path='office.docx')
print(doc)
document = Document('/Users/disturber/Downloads/lekinterkaps.doc')
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            print(cell.text)
